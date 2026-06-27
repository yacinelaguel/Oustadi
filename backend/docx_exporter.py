# =============================================================================
#  OSTADI — أستاذي
#  RTL Arabic DOCX Export Engine
#  Architect: Yacine Laguel
#  File: backend/docx_exporter.py
#
#  Responsibilities:
#    1. Build a fully editable Microsoft Word (.docx) exam document
#    2. Render the official four-zone Ministry header as a Word table
#    3. Apply RTL paragraph direction on every single paragraph
#    4. Embed Amiri font references throughout all styles
#    5. Render text passage (literary subjects), sections, questions,
#       sub-questions with correct indentation and point labels
#    6. Apply professional line spacing, borders, and shading
#
#  Dependencies:
#    pip install python-docx lxml
#
#  The output .docx opens natively in Microsoft Word and LibreOffice
#  with correct RTL layout. Teachers can edit content freely.
#
#  Called exclusively by main.py → generate_exam() → build_exam_docx()
# =============================================================================

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.table import Table, _Cell

# ---------------------------------------------------------------------------
# MODULE LOGGER
# ---------------------------------------------------------------------------
log = logging.getLogger("ostadi.docx_exporter")

# ---------------------------------------------------------------------------
# COLOUR CONSTANTS (RGB tuples for python-docx)
# ---------------------------------------------------------------------------
COLOR_NAVY        = RGBColor(0x1A, 0x3A, 0x6B)   # Ministry dark navy
COLOR_BLACK       = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK_GREY   = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_MID_GREY    = RGBColor(0x4A, 0x4A, 0x4A)
COLOR_RED_DARK    = RGBColor(0x8B, 0x00, 0x00)   # Correction headings
COLOR_LIGHT_BG    = "F4F6FA"                       # Hex string for XML shading

# ---------------------------------------------------------------------------
# FONT NAME
# ---------------------------------------------------------------------------
ARABIC_FONT = "Amiri"


# =============================================================================
#  SECTION 1: LOW-LEVEL XML HELPERS
#  python-docx exposes only a subset of OOXML features through its
#  Python API. The helpers below manipulate the underlying XML directly
#  to implement RTL direction, paragraph borders, cell shading, and
#  custom spacing — features not available through the standard API.
# =============================================================================

def _set_rtl(paragraph) -> None:
    """
    Forces Right-to-Left text direction on a paragraph.
    Sets both the paragraph-level bidi flag and the run-level rtl flag
    so that Arabic text renders correctly in all Word versions.
    """
    pPr = paragraph._p.get_or_add_pPr()

    # Paragraph-level bidi
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

    # Also set jc (justification) to right for RTL
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")


def _set_run_rtl(run) -> None:
    """
    Sets the RTL flag on a specific run's rPr (run properties).
    Required alongside paragraph-level RTL for correct glyph ordering.
    """
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def _set_para_spacing(
    paragraph,
    before_pt: float = 0,
    after_pt: float  = 0,
    line_rule: str   = "auto",
    line_val: int    = 276,
) -> None:
    """
    Sets paragraph spacing (before, after, line height) via XML.
    line_val is in twentieths of a point (twips).
    276 twips = 1.15× line spacing (comfortable Arabic reading height).
    """
    pPr   = paragraph._p.get_or_add_pPr()
    pSpc  = OxmlElement("w:spacing")
    pSpc.set(qn("w:before"), str(int(before_pt * 20)))
    pSpc.set(qn("w:after"),  str(int(after_pt  * 20)))
    pSpc.set(qn("w:line"),   str(line_val))
    pSpc.set(qn("w:lineRule"), line_rule)
    pPr.append(pSpc)


def _set_cell_shading(cell: _Cell, fill_hex: str) -> None:
    """
    Applies a solid background fill to a table cell.
    fill_hex: six-character hex string without '#', e.g. "F4F6FA"
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_cell_border(
    cell: _Cell,
    top: bool    = False,
    bottom: bool = False,
    left: bool   = False,
    right: bool  = False,
    color_hex: str = "2B4C8C",
    size_pt: int   = 8,
) -> None:
    """
    Applies selective borders to a table cell.
    size_pt is in eighth-points (8 = 1pt border).
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")

    border_map = {
        "top":    top,
        "bottom": bottom,
        "left":   left,
        "right":  right,
    }

    for side, apply_border in border_map.items():
        el = OxmlElement(f"w:{side}")
        if apply_border:
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    str(size_pt))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color_hex)
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)

    tcPr.append(tcBorders)


def _remove_table_borders(table: Table) -> None:
    """
    Removes all visible borders from a table — used for layout tables
    that should be invisible (header layout, student info bar).
    """
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)

    tblPr.append(tblBorders)


def _set_table_width(table: Table, width_cm: float) -> None:
    """
    Sets the preferred width of a table to an exact cm value.
    """
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(int(width_cm * 567)))   # 1cm = 567 twips
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _add_horizontal_rule(document: Document, color_hex: str = "2B4C8C") -> None:
    """
    Adds a full-width horizontal rule paragraph using paragraph borders.
    Simulates an <hr> element in the Word document.
    """
    p   = document.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    pBdr = OxmlElement("w:pBdr")
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"),   "single")
    bottom_border.set(qn("w:sz"),    "12")       # 1.5pt
    bottom_border.set(qn("w:space"), "1")
    bottom_border.set(qn("w:color"), color_hex)
    pBdr.append(bottom_border)
    pPr.append(pBdr)

    _set_para_spacing(p, before_pt=2, after_pt=2)
    return p


# =============================================================================
#  SECTION 2: PARAGRAPH & RUN FACTORY
#  Centralised helpers that create styled paragraphs and runs
#  with RTL direction and Amiri font pre-applied.
# =============================================================================

def _add_arabic_paragraph(
    document: Document,
    text: str,
    font_size_pt: float     = 12,
    bold: bool              = False,
    italic: bool            = False,
    color: RGBColor         = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT,
    space_before_pt: float  = 0,
    space_after_pt: float   = 4,
    line_val: int           = 276,
    indent_right_cm: float  = 0,
    indent_left_cm: float   = 0,
) -> Any:
    """
    Adds a single RTL Arabic paragraph to the document with full styling.

    Args:
        document:        The Document object to append to.
        text:            Raw Arabic text (NOT pre-shaped — DOCX handles shaping).
        font_size_pt:    Font size in points.
        bold:            Bold weight flag.
        italic:          Italic flag.
        color:           RGBColor instance or None (defaults to COLOR_DARK_GREY).
        alignment:       WD_ALIGN_PARAGRAPH constant.
        space_before_pt: Paragraph spacing before in points.
        space_after_pt:  Paragraph spacing after in points.
        line_val:        Line height in twips (276 = 1.15×).
        indent_right_cm: Right indent in cm.
        indent_left_cm:  Left indent in cm.

    Returns:
        The created paragraph object.

    NOTE: Unlike PDF export, DOCX does NOT require arabic_reshaper.
    Microsoft Word handles Arabic glyph shaping natively.
    The RTL flag is applied at both paragraph and run level.
    """
    if color is None:
        color = COLOR_DARK_GREY

    paragraph = document.add_paragraph()
    paragraph.alignment = alignment

    # Apply RTL at paragraph level
    _set_rtl(paragraph)
    _set_para_spacing(
        paragraph,
        before_pt=space_before_pt,
        after_pt=space_after_pt,
        line_val=line_val,
    )

    # Apply indentation if requested
    if indent_right_cm > 0 or indent_left_cm > 0:
        pPr = paragraph._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        if indent_right_cm > 0:
            ind.set(qn("w:right"), str(int(indent_right_cm * 567)))
        if indent_left_cm > 0:
            ind.set(qn("w:left"),  str(int(indent_left_cm  * 567)))
        pPr.append(ind)

    # Create and style the run
    run = paragraph.add_run(text)
    run.font.name  = ARABIC_FONT
    run.font.size  = Pt(font_size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color

    # Apply RTL at run level
    _set_run_rtl(run)

    # Force Amiri for complex scripts (Arabic)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),       ARABIC_FONT)
    rFonts.set(qn("w:hAnsi"),       ARABIC_FONT)
    rFonts.set(qn("w:cs"),          ARABIC_FONT)
    rFonts.set(qn("w:eastAsia"),    ARABIC_FONT)

    return paragraph


def _add_arabic_paragraph_to_cell(
    cell: _Cell,
    text: str,
    font_size_pt: float     = 11,
    bold: bool              = False,
    color: RGBColor         = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT,
) -> Any:
    """
    Adds an RTL Arabic paragraph inside a table cell.
    Clears any default empty paragraph that python-docx inserts automatically.
    """
    if color is None:
        color = COLOR_DARK_GREY

    # Clear default empty paragraph
    for existing_para in cell.paragraphs:
        existing_para._p.getparent().remove(existing_para._p)

    paragraph = cell.add_paragraph()
    paragraph.alignment = alignment
    _set_rtl(paragraph)
    _set_para_spacing(paragraph, before_pt=2, after_pt=2, line_val=240)

    run = paragraph.add_run(text)
    run.font.name      = ARABIC_FONT
    run.font.size      = Pt(font_size_pt)
    run.font.bold      = bold
    if color:
        run.font.color.rgb = color
    _set_run_rtl(run)

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    ARABIC_FONT)
    rFonts.set(qn("w:hAnsi"),    ARABIC_FONT)
    rFonts.set(qn("w:cs"),       ARABIC_FONT)
    rFonts.set(qn("w:eastAsia"), ARABIC_FONT)

    return paragraph


# =============================================================================
#  SECTION 3: DOCUMENT SETUP
#  Configures the document-level settings:
#    - A4 page size
#    - RTL document direction
#    - Margins
#    - Default font
# =============================================================================

def _setup_document() -> Document:
    """
    Creates and returns a new Document configured for RTL Arabic output.
    Sets A4 page size, professional margins, and document-level BiDi.
    """
    document = Document()

    # ---- Page size: A4 ----
    section = document.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    # ---- Document-level RTL direction ----
    settings   = document.settings.element
    doc_defaults = settings.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        settings.append(doc_defaults)

    # Set document bidi direction to RTL
    bidi_default = OxmlElement("w:bidi")
    bidi_default.set(qn("w:val"), "1")

    rPrDefault = doc_defaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        doc_defaults.append(rPrDefault)

    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)

    # Default complex script font = Amiri
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    ARABIC_FONT)
    rFonts.set(qn("w:hAnsi"),    ARABIC_FONT)
    rFonts.set(qn("w:cs"),       ARABIC_FONT)
    rFonts.set(qn("w:eastAsia"), ARABIC_FONT)
    rPr.append(rFonts)

    log.debug("Document created: A4, RTL, Amiri default font")
    return document


# =============================================================================
#  SECTION 4: MINISTRY HEADER BUILDER
# =============================================================================

def _build_docx_ministry_header(
    document: Document,
    academic_year: str,
    specialty_stream: str,
    subject: str,
    trimester: str,
    duration: int,
    is_correction: bool = False,
) -> None:
    """
    Inserts the official four-zone Algerian Ministry exam header
    into the document as invisible layout tables.

    Zone 1 (Top Row):
      Right cell → Republic + Ministry labels (bold navy)
      Left cell  → Wilaya + School dotted fields

    Zone 2: Full-width navy horizontal rule

    Zone 3: Centred exam title block

    Zone 4: Student info bar (exam only)

    Zone 5: Second horizontal rule
    """

    # ---- ZONE 1: TOP TWO-COLUMN HEADER TABLE ----
    header_table = document.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(header_table)
    _set_table_width(header_table, 17.0)   # 17cm content width

    # Column widths: each column ~8.5cm
    for cell in header_table.rows[0].cells:
        cell.width = Cm(8.5)

    right_cell = header_table.rows[0].cells[1]   # RTL: right is index 1
    left_cell  = header_table.rows[0].cells[0]   # RTL: left  is index 0

    # Right cell — Republic + Ministry
    _add_arabic_paragraph_to_cell(
        right_cell,
        "الجمهورية الجزائرية الديمقراطية الشعبية",
        font_size_pt=10,
        bold=True,
        color=COLOR_NAVY,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    _add_arabic_paragraph_to_cell(
        right_cell,
        "وزارة التربية الوطنية",
        font_size_pt=10,
        bold=True,
        color=COLOR_NAVY,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    # Left cell — Wilaya + School fields
    _add_arabic_paragraph_to_cell(
        left_cell,
        "مديرية التربية لولاية: ................................",
        font_size_pt=10,
        bold=False,
        color=COLOR_DARK_GREY,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _add_arabic_paragraph_to_cell(
        left_cell,
        "مؤسسة: ................................................",
        font_size_pt=10,
        bold=False,
        color=COLOR_DARK_GREY,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    # ---- ZONE 2: FIRST HORIZONTAL RULE ----
    _add_horizontal_rule(document, color_hex="2B4C8C")

    # ---- ZONE 3: CENTRE TITLE BLOCK ----
    doc_type = "التصحيح النموذجي وسلّم التنقيط" if is_correction else "اختبار الثلاثي"
    title_text = f"{doc_type} ({trimester}) — السنة الدراسية: 2025–2026م"

    _add_arabic_paragraph(
        document,
        title_text,
        font_size_pt=14,
        bold=True,
        color=COLOR_NAVY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=4,
        space_after_pt=2,
        line_val=300,
    )

    # Exam details line
    details_text = (
        f"المادة: {subject}    |    "
        f"الشعبة: {specialty_stream}    |    "
        f"المدة: {duration} دقيقة"
    )
    _add_arabic_paragraph(
        document,
        details_text,
        font_size_pt=10,
        bold=False,
        color=COLOR_MID_GREY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=0,
        space_after_pt=4,
        line_val=240,
    )

    # ---- ZONE 4: SECOND HORIZONTAL RULE ----
    _add_horizontal_rule(document, color_hex="2B4C8C")

    # ---- ZONE 5: STUDENT INFO BAR (exam only) ----
    if not is_correction:
        student_table = document.add_table(rows=1, cols=3)
        student_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(student_table)
        _set_table_width(student_table, 17.0)

        student_labels = [
            "القسم: ............................",
            "الاسم: ............................",
            "اللقب: ............................",
        ]

        for idx, cell in enumerate(student_table.rows[0].cells):
            cell.width = Cm(17.0 / 3)
            _add_arabic_paragraph_to_cell(
                cell,
                student_labels[idx],
                font_size_pt=10,
                bold=False,
                color=COLOR_DARK_GREY,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )

        # Add bottom border to student table cells
        for cell in student_table.rows[0].cells:
            _set_cell_border(cell, bottom=True, color_hex="CCCCCC", size_pt=4)

        # Spacer paragraph after student bar
        spacer = document.add_paragraph()
        _set_para_spacing(spacer, before_pt=0, after_pt=6)


# =============================================================================
#  SECTION 5: EXAM BODY RENDERER
# =============================================================================

def _render_docx_exam_body(
    document: Document,
    exam_obj: Dict[str, Any],
) -> None:
    """
    Renders the full exam body into the document.
    Processes text passage (literary), section headers,
    questions, and sub-questions with point labels.
    """

    # ---- TEXT PASSAGE ----
    text_passage = exam_obj.get("text_passage")
    if text_passage and isinstance(text_passage, str) and text_passage.strip():

        _add_arabic_paragraph(
            document,
            "النص:",
            font_size_pt=12,
            bold=True,
            color=COLOR_NAVY,
            space_before_pt=6,
            space_after_pt=2,
        )

        # Shaded passage box — use a single-cell table for background
        passage_table = document.add_table(rows=1, cols=1)
        passage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_width(passage_table, 17.0)
        passage_cell = passage_table.rows[0].cells[0]
        _set_cell_shading(passage_cell, COLOR_LIGHT_BG)
        _set_cell_border(
            passage_cell,
            top=True, bottom=True, left=True, right=True,
            color_hex="CCCCCC",
            size_pt=4,
        )

        for line in text_passage.strip().split("\n"):
            stripped = line.strip()
            if stripped:
                _add_arabic_paragraph_to_cell(
                    passage_cell,
                    stripped,
                    font_size_pt=12,
                    bold=False,
                    color=COLOR_DARK_GREY,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                )

        # Spacer after passage
        spacer = document.add_paragraph()
        _set_para_spacing(spacer, before_pt=0, after_pt=4)

    # ---- SECTIONS ----
    sections: List[Dict] = exam_obj.get("sections", [])

    for section in sections:
        section_title = section.get("section_title", "")
        points_total  = section.get("points_total", "")
        questions     = section.get("questions", [])

        # Section header: two-column table (title right | points left)
        sec_table = document.add_table(rows=1, cols=2)
        sec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(sec_table)
        _set_table_width(sec_table, 17.0)

        title_cell  = sec_table.rows[0].cells[1]   # RTL: right
        points_cell = sec_table.rows[0].cells[0]   # RTL: left

        title_cell.width  = Cm(13.0)
        points_cell.width = Cm(4.0)

        _add_arabic_paragraph_to_cell(
            title_cell,
            section_title,
            font_size_pt=12,
            bold=True,
            color=COLOR_NAVY,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        points_label = f"({points_total} نقاط)" if points_total else ""
        _add_arabic_paragraph_to_cell(
            points_cell,
            points_label,
            font_size_pt=11,
            bold=True,
            color=COLOR_NAVY,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )

        # Bottom border on section header row
        for cell in sec_table.rows[0].cells:
            _set_cell_border(cell, bottom=True, color_hex="2B4C8C", size_pt=8)

        # Spacer
        spacer = document.add_paragraph()
        _set_para_spacing(spacer, before_pt=0, after_pt=3)

        # ---- QUESTIONS ----
        for question in questions:
            q_number      = question.get("number", "")
            q_text        = question.get("text", "")
            q_points      = question.get("points")
            sub_questions = question.get("sub_questions", [])

            q_label = f"{q_number}—  {q_text}" if q_number else q_text

            if q_points is not None and not sub_questions:
                # Direct points — two column layout
                q_table = document.add_table(rows=1, cols=2)
                q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                _remove_table_borders(q_table)
                _set_table_width(q_table, 17.0)

                q_text_cell   = q_table.rows[0].cells[1]
                q_points_cell = q_table.rows[0].cells[0]

                q_text_cell.width   = Cm(13.5)
                q_points_cell.width = Cm(3.5)

                _add_arabic_paragraph_to_cell(
                    q_text_cell,
                    q_label,
                    font_size_pt=11,
                    bold=True,
                    color=COLOR_DARK_GREY,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                )
                _add_arabic_paragraph_to_cell(
                    q_points_cell,
                    f"({q_points} ن)",
                    font_size_pt=10,
                    bold=True,
                    color=COLOR_NAVY,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                )

            else:
                # No direct points — full width question label
                _add_arabic_paragraph(
                    document,
                    q_label,
                    font_size_pt=11,
                    bold=True,
                    color=COLOR_DARK_GREY,
                    space_before_pt=4,
                    space_after_pt=2,
                )

            # ---- SUB-QUESTIONS ----
            for sub in sub_questions:
                sub_number = sub.get("number", "")
                sub_text   = sub.get("text", "")
                sub_points = sub.get("points", "")

                sub_label = f"{sub_number})  {sub_text}" if sub_number else sub_text

                sub_table = document.add_table(rows=1, cols=2)
                sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                _remove_table_borders(sub_table)
                _set_table_width(sub_table, 17.0)

                sub_text_cell   = sub_table.rows[0].cells[1]
                sub_points_cell = sub_table.rows[0].cells[0]

                sub_text_cell.width   = Cm(13.5)
                sub_points_cell.width = Cm(3.5)

                _add_arabic_paragraph_to_cell(
                    sub_text_cell,
                    sub_label,
                    font_size_pt=11,
                    bold=False,
                    color=COLOR_DARK_GREY,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                )
                _add_arabic_paragraph_to_cell(
                    sub_points_cell,
                    f"({sub_points} ن)" if sub_points else "",
                    font_size_pt=10,
                    bold=True,
                    color=COLOR_NAVY,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                )

        # Section spacer
        spacer = document.add_paragraph()
        _set_para_spacing(spacer, before_pt=0, after_pt=6)


# =============================================================================
#  SECTION 6: PUBLIC API — DOCX BUILDER
# =============================================================================

def build_exam_docx(
    output_path: Path,
    payload: Dict[str, Any],
    academic_year: str,
    specialty_stream: str,
    subject: str,
    trimester: str,
    duration: int,
    fonts_dir: Path,
) -> None:
    """
    Compiles the official editable exam DOCX from the Gemini payload.

    The resulting file:
      - Opens natively in Microsoft Word and LibreOffice Writer
      - Renders all Arabic text in Amiri with correct RTL direction
      - Contains the full Ministry header, exam body, sections,
        questions, sub-questions, and point labels
      - Is fully editable by the teacher

    Args:
        output_path:      Absolute path where the .docx will be saved.
        payload:          Validated Gemini JSON payload dict.
        academic_year:    e.g. "2AS"
        specialty_stream: e.g. "علوم تجريبية"
        subject:          e.g. "الرياضيات"
        trimester:        e.g. "الفصل الأول"
        duration:         Exam duration in minutes.
        fonts_dir:        Path to backend/fonts/ — not used directly by
                          python-docx but kept for API consistency.
    """
    document = _setup_document()

    # ---- MINISTRY HEADER ----
    _build_docx_ministry_header(
        document=document,
        academic_year=academic_year,
        specialty_stream=specialty_stream,
        subject=subject,
        trimester=trimester,
        duration=duration,
        is_correction=False,
    )

    # ---- EXAM BODY ----
    exam_obj = payload.get("exam", {})
    _render_docx_exam_body(document=document, exam_obj=exam_obj)

    # ---- SAVE ----
    document.save(str(output_path))
    log.info(
        f"Exam DOCX written → {output_path.name} "
        f"({output_path.stat().st_size / 1024:.1f} KB)"
    )